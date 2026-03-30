# -*- coding: utf-8 -*-
"""マニュアルMarkdownファイルをWord(.docx)に変換するスクリプト"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

DOCS_DIR = Path(__file__).parent
MD_FILES = [
    ("01_担任用簡易マニュアル.md", "担任用簡易マニュアル"),
    ("02_詳細マニュアル.md", "詳細マニュアル"),
    ("03_管理マニュアル.md", "管理マニュアル"),
    ("04_引き継ぎマニュアル.md", "引き継ぎマニュアル"),
]


def set_run_font(run, size=10.5, bold=False, color=None, name="游ゴシック"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_text(doc, text, level=1):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    sizes = {1: 16, 2: 14, 3: 12}
    set_run_font(run, size=sizes.get(level, 12), bold=True)


def add_paragraph_text(doc, text, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    if italic:
        run.font.italic = True
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ヘッダー行
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=9, bold=True)

    # データ行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_run_font(run, size=9)

    return table


def parse_table_block(lines):
    """Markdownのテーブル行をパースしてheaders, rowsを返す"""
    headers = [c.strip() for c in lines[0].strip("|").split("|")]
    rows = []
    for line in lines[2:]:  # 2行目はセパレータ
        cols = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cols)
    return headers, rows


def convert_md_to_docx(md_path, title, output_path):
    doc = Document()

    # ページ設定（A4縦）
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # デフォルトフォント
    style = doc.styles["Normal"]
    style.font.name = "游ゴシック"
    style.font.size = Pt(10.5)

    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 空行
        if not stripped:
            i += 1
            continue

        # 水平線
        if stripped in ("---", "***", "___"):
            # doc.add_paragraph("─" * 50)
            i += 1
            continue

        # 見出し
        hm = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if hm:
            level = len(hm.group(1))
            add_heading_text(doc, hm.group(2), level=min(level, 3))
            i += 1
            continue

        # テーブル
        if "|" in stripped and i + 1 < len(lines) and re.match(
            r"^\|[\s\-:|]+\|$", lines[i + 1].strip()
        ):
            table_lines = []
            while i < len(lines) and "|" in lines[i].strip():
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 3:
                headers, rows = parse_table_block(table_lines)
                add_table(doc, headers, rows)
            continue

        # 引用
        if stripped.startswith(">"):
            text_content = stripped.lstrip("> ").strip()
            p = add_paragraph_text(doc, text_content, italic=True, size=9.5)
            p.paragraph_format.left_indent = Cm(0.5)
            i += 1
            continue

        # コードブロック
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # 閉じ```をスキップ
            code_text = "\n".join(code_lines)
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            set_run_font(run, size=8.5, name="Consolas")
            p.paragraph_format.left_indent = Cm(0.5)
            continue

        # リスト
        lm = re.match(r"^(\s*)[-*]\s+(.+)$", stripped)
        if lm:
            text_content = lm.group(2)
            # **太字** を処理
            p = doc.add_paragraph(style="List Bullet")
            parts = re.split(r"(\*\*[^*]+\*\*)", text_content)
            for part in parts:
                bm = re.match(r"^\*\*(.+)\*\*$", part)
                if bm:
                    run = p.add_run(bm.group(1))
                    set_run_font(run, size=10, bold=True)
                else:
                    run = p.add_run(part)
                    set_run_font(run, size=10)
            i += 1
            continue

        # 番号付きリスト
        nm = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if nm:
            p = doc.add_paragraph(style="List Number")
            text_content = nm.group(2)
            parts = re.split(r"(\*\*[^*]+\*\*)", text_content)
            for part in parts:
                bm = re.match(r"^\*\*(.+)\*\*$", part)
                if bm:
                    run = p.add_run(bm.group(1))
                    set_run_font(run, size=10, bold=True)
                else:
                    run = p.add_run(part)
                    set_run_font(run, size=10)
            i += 1
            continue

        # 通常段落（**太字**対応）
        p = doc.add_paragraph()
        parts = re.split(r"(\*\*[^*]+\*\*)", stripped)
        for part in parts:
            bm = re.match(r"^\*\*(.+)\*\*$", part)
            if bm:
                run = p.add_run(bm.group(1))
                set_run_font(run, size=10.5, bold=True)
            else:
                # `コード` 対応
                code_parts = re.split(r"(`[^`]+`)", part)
                for cp in code_parts:
                    cm = re.match(r"^`(.+)`$", cp)
                    if cm:
                        run = p.add_run(cm.group(1))
                        set_run_font(run, size=9.5, name="Consolas")
                    else:
                        run = p.add_run(cp)
                        set_run_font(run, size=10.5)
        i += 1

    doc.save(str(output_path))
    return output_path


def main():
    output_dir = DOCS_DIR / "word"
    output_dir.mkdir(exist_ok=True)

    for md_name, title in MD_FILES:
        md_path = DOCS_DIR / md_name
        if not md_path.exists():
            print(f"  スキップ: {md_name} が見つかりません")
            continue
        out_path = output_dir / md_name.replace(".md", ".docx")
        convert_md_to_docx(md_path, title, out_path)
        print(f"  生成: {out_path.name}")

    print(f"\n完了: {output_dir}/")


if __name__ == "__main__":
    main()
